"""
Document processing service for handling file uploads and content extraction.
"""
import os
import shutil
import magic
from pathlib import Path
from typing import Optional, Tuple, Dict, Any
import structlog
import PyPDF2
import docx
from bs4 import BeautifulSoup
import hashlib
from datetime import datetime

from app.core.config import settings
from app.core.exceptions import DocumentProcessingException

logger = structlog.get_logger()


class DocumentProcessor:
    """Service for processing and extracting content from documents."""
    
    def __init__(self):
        self.upload_dir = Path(settings.UPLOAD_DIRECTORY)
        self.upload_dir.mkdir(exist_ok=True)
        self.allowed_types = settings.ALLOWED_FILE_TYPES
        self.max_size = settings.MAX_UPLOAD_SIZE
    
    def validate_file(self, file_content: bytes, filename: str) -> Tuple[bool, str]:
        """
        Validate uploaded file.
        
        Args:
            file_content: Binary content of the file
            filename: Original filename
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        try:
            # Check file size
            if len(file_content) > self.max_size:
                return False, f"File size exceeds maximum allowed size of {self.max_size} bytes"
            
            # Check file extension
            file_ext = Path(filename).suffix.lower().lstrip('.')
            if file_ext not in self.allowed_types:
                return False, f"File type '{file_ext}' is not allowed. Allowed types: {', '.join(self.allowed_types)}"
            
            # Check MIME type
            try:
                mime_type = magic.from_buffer(file_content, mime=True)
                if not self._is_mime_type_allowed(mime_type, file_ext):
                    return False, f"MIME type '{mime_type}' does not match file extension '{file_ext}'"
            except Exception as e:
                logger.warning(f"Could not determine MIME type: {e}")
                # Continue without MIME type validation
            
            return True, ""
            
        except Exception as e:
            logger.error(f"Error validating file: {e}")
            return False, f"File validation error: {e}"
    
    def _is_mime_type_allowed(self, mime_type: str, file_ext: str) -> bool:
        """Check if MIME type matches file extension."""
        mime_mappings = {
            'pdf': ['application/pdf'],
            'docx': ['application/vnd.openxmlformats-officedocument.wordprocessingml.document'],
            'txt': ['text/plain'],
            'html': ['text/html', 'application/html'],
        }
        
        return mime_type in mime_mappings.get(file_ext, [])
    
    def save_file(self, file_content: bytes, filename: str, user_id: int) -> Tuple[str, str]:
        """
        Save file to disk.
        
        Args:
            file_content: Binary content of the file
            filename: Original filename
            user_id: ID of the user uploading the file
            
        Returns:
            Tuple of (file_path, generated_filename)
        """
        try:
            # Generate unique filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            file_hash = hashlib.md5(file_content).hexdigest()[:8]
            file_ext = Path(filename).suffix.lower()
            generated_filename = f"{user_id}_{timestamp}_{file_hash}{file_ext}"
            
            # Create user directory
            user_dir = self.upload_dir / str(user_id)
            user_dir.mkdir(exist_ok=True)
            
            # Save file
            file_path = user_dir / generated_filename
            with open(file_path, 'wb') as f:
                f.write(file_content)
            
            logger.info(f"File saved: {file_path}")
            return str(file_path), generated_filename
            
        except Exception as e:
            logger.error(f"Error saving file: {e}")
            raise DocumentProcessingException(f"Failed to save file: {e}")
    
    def extract_content(self, file_path: str, file_type: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text content and metadata from file.
        
        Args:
            file_path: Path to the file
            file_type: Type of the file (pdf, docx, txt, html)
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        try:
            if file_type == 'pdf':
                return self._extract_pdf_content(file_path)
            elif file_type == 'docx':
                return self._extract_docx_content(file_path)
            elif file_type == 'txt':
                return self._extract_txt_content(file_path)
            elif file_type == 'html':
                return self._extract_html_content(file_path)
            else:
                raise DocumentProcessingException(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            logger.error(f"Error extracting content from {file_path}: {e}")
            raise DocumentProcessingException(f"Failed to extract content: {e}")
    
    def _extract_pdf_content(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract content from PDF file."""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            
            metadata = {
                "page_count": len(reader.pages),
                "pdf_info": reader.metadata,
            }
            
            return text.strip(), metadata
    
    def _extract_docx_content(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract content from DOCX file."""
        doc = docx.Document(file_path)
        
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        metadata = {
            "paragraph_count": len(doc.paragraphs),
            "core_properties": {
                "title": doc.core_properties.title,
                "author": doc.core_properties.author,
                "subject": doc.core_properties.subject,
                "created": str(doc.core_properties.created) if doc.core_properties.created else None,
                "modified": str(doc.core_properties.modified) if doc.core_properties.modified else None,
            }
        }
        
        return text.strip(), metadata
    
    def _extract_txt_content(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract content from TXT file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        metadata = {
            "line_count": len(text.splitlines()),
            "character_count": len(text),
        }
        
        return text.strip(), metadata
    
    def _extract_html_content(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract content from HTML file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            html_content = file.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Extract text content
        text = soup.get_text(separator='\n', strip=True)
        
        # Extract metadata
        title = soup.find('title')
        meta_description = soup.find('meta', attrs={'name': 'description'})
        meta_keywords = soup.find('meta', attrs={'name': 'keywords'})
        
        metadata = {
            "title": title.get_text() if title else None,
            "description": meta_description.get('content') if meta_description else None,
            "keywords": meta_keywords.get('content') if meta_keywords else None,
        }
        
        return text.strip(), metadata
    
    def get_content_preview(self, content: str, max_length: int = 500) -> str:
        """
        Generate a preview of the content.
        
        Args:
            content: Full text content
            max_length: Maximum length of preview
            
        Returns:
            Content preview
        """
        if len(content) <= max_length:
            return content
        
        # Find the last complete sentence within the limit
        preview = content[:max_length]
        last_sentence_end = max(
            preview.rfind('.'),
            preview.rfind('!'),
            preview.rfind('?')
        )
        
        if last_sentence_end > max_length * 0.7:  # At least 70% of max_length
            preview = preview[:last_sentence_end + 1]
        else:
            preview = preview + "..."
        
        return preview
    
    def delete_file(self, file_path: str) -> bool:
        """
        Delete file from disk.
        
        Args:
            file_path: Path to the file
            
        Returns:
            True if successful
        """
        try:
            path = Path(file_path)
            if path.exists():
                path.unlink()
                logger.info(f"File deleted: {file_path}")
                return True
            else:
                logger.warning(f"File not found for deletion: {file_path}")
                return False
                
        except Exception as e:
            logger.error(f"Error deleting file {file_path}: {e}")
            return False
    
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """
        Get file information.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary with file information
        """
        try:
            path = Path(file_path)
            stat = path.stat()
            
            return {
                "size": stat.st_size,
                "created": datetime.fromtimestamp(stat.st_ctime),
                "modified": datetime.fromtimestamp(stat.st_mtime),
                "exists": path.exists(),
            }
            
        except Exception as e:
            logger.error(f"Error getting file info for {file_path}: {e}")
            return {"exists": False, "error": str(e)}


# Global processor instance
document_processor = DocumentProcessor()
